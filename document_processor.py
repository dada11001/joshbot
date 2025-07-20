import PyPDF2
import docx
from docx import Document
import streamlit as st
import os
from typing import Optional
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None
from PIL import Image
import pytesseract
import io
import base64

class DocumentProcessor:
    """
    Handles document processing for PDF and DOC/DOCX files
    Extracts text content for further processing
    """
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.doc', '.docx']
        self.use_ocr = True  # Enable OCR for image processing
    
    def extract_text(self, file_path: str) -> str:
        """
        Extract text from uploaded document based on file extension
        
        Args:
            file_path (str): Path to the uploaded file
            
        Returns:
            str: Extracted text content
        """
        try:
            file_extension = os.path.splitext(file_path)[1].lower()
            
            if file_extension == '.pdf':
                return self._extract_from_pdf_advanced(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from document: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file using PyPDF2
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text content
        """
        text_content = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise Exception("PDF is password-protected. Please provide an unprotected version.")
                
                # Extract text from each page
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text + "\n"
                    except Exception as page_error:
                        st.warning(f"Could not extract text from page {page_num + 1}: {str(page_error)}")
                        continue
                
                if not text_content.strip():
                    raise Exception("No readable text found in PDF. The document might be image-based or corrupted.")
                
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
        
        return self._clean_text(text_content)
    
    def _extract_from_pdf_advanced(self, file_path: str) -> str:
        """
        Advanced PDF text extraction using PyMuPDF with OCR for images
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            str: Extracted text content including text from images
        """
        text_content = ""
        
        # Check if PyMuPDF is available
        if not fitz:
            st.warning("PyMuPDF not available. Using basic PDF extraction.")
            return self._extract_from_pdf(file_path)
        
        try:
            # Try PyMuPDF first for better text extraction
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract regular text
                page_text = page.get_text("text")
                
                if page_text.strip():
                    text_content += f"\n--- Page {page_num + 1} ---\n"
                    text_content += page_text + "\n"
                
                # Extract text from images on the page if OCR is enabled
                if self.use_ocr:
                    try:
                        image_text = self._extract_text_from_images_on_page(page, page_num + 1)
                        if image_text.strip():
                            text_content += f"\n--- Images from Page {page_num + 1} ---\n"
                            text_content += image_text + "\n"
                    except Exception as img_error:
                        st.warning(f"Could not process images on page {page_num + 1}: {str(img_error)}")
                        continue
            
            doc.close()
            
            # If no text was extracted, fall back to PyPDF2
            if not text_content.strip():
                st.info("Using fallback text extraction method...")
                return self._extract_from_pdf(file_path)
                
        except Exception as e:
            st.warning(f"Advanced PDF processing failed: {str(e)}. Using fallback method...")
            return self._extract_from_pdf(file_path)
        
        return self._clean_text(text_content)
    
    def _extract_text_from_images_on_page(self, page, page_num: int) -> str:
        """
        Extract text from images on a PDF page using OCR
        
        Args:
            page: PyMuPDF page object
            page_num (int): Page number for logging
            
        Returns:
            str: Extracted text from images
        """
        image_text = ""
        
        # Check if PyMuPDF is available
        if not fitz:
            return ""
        
        try:
            # Get images from the page
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    # Get image data
                    xref = img[0]
                    pix = fitz.Pixmap(page.parent, xref)
                    
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        # Convert to PIL Image
                        img_data = pix.tobytes("png")
                        pil_image = Image.open(io.BytesIO(img_data))
                        
                        # Use OCR to extract text
                        ocr_text = pytesseract.image_to_string(pil_image, lang='eng')
                        
                        if ocr_text.strip():
                            image_text += f"\n[Image {img_index + 1} on page {page_num}]\n"
                            image_text += ocr_text + "\n"
                    
                    pix = None  # Free memory
                    
                except Exception as img_error:
                    st.warning(f"Could not process image {img_index + 1} on page {page_num}: {str(img_error)}")
                    continue
                    
        except Exception as e:
            st.warning(f"Error processing images on page {page_num}: {str(e)}")
        
        return image_text
    
    def _extract_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOC/DOCX file using python-docx
        
        Args:
            file_path (str): Path to DOC/DOCX file
            
        Returns:
            str: Extracted text content
        """
        text_content = ""
        
        try:
            # Try to read as DOCX first
            doc = Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + " "
                    text_content += "\n"
            
            if not text_content.strip():
                raise Exception("No readable text found in document.")
                
        except Exception as e:
            # If it's a .doc file, it might not be supported by python-docx
            if file_path.lower().endswith('.doc'):
                raise Exception("Legacy .doc files are not fully supported. Please convert to .docx format.")
            else:
                raise Exception(f"Error reading document: {str(e)}")
        
        return self._clean_text(text_content)
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text (str): Raw extracted text
            
        Returns:
            str: Cleaned text content
        """
        if not text:
            return ""
        
        # Remove excessive whitespace and normalize line breaks
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            cleaned_line = line.strip()
            if cleaned_line:  # Skip empty lines
                cleaned_lines.append(cleaned_line)
        
        # Join lines with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove excessive spaces
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        return cleaned_text
    
    def validate_file(self, file_path: str) -> bool:
        """
        Validate if file exists and has supported extension
        
        Args:
            file_path (str): Path to file
            
        Returns:
            bool: True if file is valid, False otherwise
        """
        if not os.path.exists(file_path):
            return False
        
        file_extension = os.path.splitext(file_path)[1].lower()
        return file_extension in self.supported_extensions
    
    def get_file_info(self, file_path: str) -> dict:
        """
        Get basic information about the uploaded file
        
        Args:
            file_path (str): Path to file
            
        Returns:
            dict: File information
        """
        if not os.path.exists(file_path):
            return {}
        
        file_stats = os.stat(file_path)
        file_extension = os.path.splitext(file_path)[1].lower()
        
        return {
            'name': os.path.basename(file_path),
            'size': file_stats.st_size,
            'extension': file_extension,
            'supported': file_extension in self.supported_extensions
        }
